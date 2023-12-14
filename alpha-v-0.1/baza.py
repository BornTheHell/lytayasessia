from docx import Document
import re
import os

def search_and_copy_questions(questions, document_path, output_path):
    # Загрузка документа Word
    doc = Document(document_path)

    # Создание нового документа для записи результатов
    output_doc = Document()

    for question in questions:
        # Используем регулярное выражение для поиска вопроса и ответов
        pattern = re.compile(f"{question}\.([^\d]+)(\d+\.)", re.DOTALL)
        match = re.search(pattern, doc.text)

        if match:
            # Записываем вопрос и правильный ответ в новый документ
            output_doc.add_paragraph(f"{question}. {match.group(1)} {match.group(2)}")
        else:
            # Если вопрос не найден, добавляем сообщение об ошибке
            output_doc.add_paragraph(f"Вопрос {question} не найден")

    # Сохранение результатов в новый документ
    output_doc.save(output_path)
    print(f"Результаты сохранены в {output_path}")

# Директория, где находится файл с вопросами
directory_path = r"C:\capcha\alpha-v-0.1"

# Получаем список файлов в директории
files_in_directory = os.listdir(directory_path)

# Выводим список файлов
print(files_in_directory)

# Проверяем, что файл "baza.doc" присутствует в списке
if "baza.doc" in files_in_directory:
    # Полный путь к файлу с вопросами
    document_path = os.path.join(directory_path, "baza.doc")

    # Путь к файлу, куда будут сохранены результаты
    output_path = "результаты.docx"

    # Список вопросов для поиска
    questions_to_search = [
        "Вопрос 1",
        "Вопрос 2",
        # ... добавьте остальные вопросы
        "Вопрос 90",
    ]

    # Запускаем функцию поиска и копирования
    search_and_copy_questions(questions_to_search, document_path, output_path)

else:
    print("Файл 'baza.doc' не найден в указанной директории.")
